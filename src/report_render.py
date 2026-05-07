from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .utils_common import _row_join_key


def _safe_float(value: Any) -> float | None:
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _format_decimal(value: float) -> str:
    return f"{value:.2f}"


def _normalize_decimal_places(text: str) -> str:
    def repl(match: re.Match[str]) -> str:
        raw = match.group(0)
        sign = "+" if raw.startswith("+") else ""
        return f"{sign}{_format_decimal(float(raw))}"

    return re.sub(r"(?<![A-Za-z0-9])[-+]?\d+\.\d+", repl, text)


def _apply_run_font(run: Any, size: float | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    if size is not None:
        run.font.size = Pt(size)


def _use_wan_yuan_display(asset: dict[str, Any]) -> bool:
    asset_name = str(asset.get("asset_name", "")).strip()
    unit = str(asset.get("unit", "")).strip()
    return asset_name == "沪深300" and unit == "亿元"


def _format_wan_yuan(value: float) -> str:
    converted = value / 10000
    text = _format_decimal(converted)
    return f"{text}万亿元"


def _normalize_hs300_units(text: str, asset: dict[str, Any]) -> str:
    if not text or not _use_wan_yuan_display(asset):
        return text

    def repl(match: re.Match[str]) -> str:
        raw = match.group(1).replace(",", "")
        try:
            return _format_wan_yuan(float(raw))
        except ValueError:
            return match.group(0)

    return re.sub(r"([+-]?\d[\d,]*(?:\.\d+)?)\s*亿元", repl, text)


def _driver_analysis_fallback(asset: dict[str, Any]) -> str:
    asset_name = str(asset.get("asset_name", "")).strip()
    metric_name = str(asset.get("metric_name", "")).strip()
    subject = " ".join(part for part in [asset_name, metric_name] if part) or "该指标"
    return (
        f"内部数据已显示{subject}本周触发异常变动；"
        "后续重点观察相关政策、资金面、市场评论与同类资产表现是否同步变化。"
    )


def _normalize_driver_analysis_text(text: str, asset: dict[str, Any] | None = None) -> str:
    fallback = _driver_analysis_fallback(asset or {})
    replacements = (
        ("驱动因素尚不明确，未检索到直接关联的外部证据。", ""),
        ("由于缺乏明确的外部驱动证据，", "结合内部数据与外部市场背景观察，"),
        ("由于暂无外部驱动证据，", "结合内部数据与外部市场背景观察，"),
        ("由于缺乏直接的驱动因素证据，", "结合内部数据与外部市场背景观察，"),
        ("由于缺乏直接驱动因素证据，", "结合内部数据与外部市场背景观察，"),
        ("由于缺乏明确的驱动证据，", "结合内部数据与外部市场背景观察，"),
        ("未检索到直接关联的外部证据。", ""),
        ("缺乏明确的外部驱动证据", "外部背景线索有限"),
        ("暂无外部驱动证据", "外部背景线索有限"),
        ("缺乏直接的驱动因素证据", "外部背景线索有限"),
        ("缺乏直接驱动因素证据", "外部背景线索有限"),
        ("缺乏明确的驱动证据", "外部背景线索有限"),
        ("驱动因素信息不足，暂作保守观察。", fallback),
        ("信息不足，保守解释。", fallback),
        ("驱动力尚不明确；", "仍需结合市场背景继续观察；"),
        ("驱动力尚不明确。", "仍需结合市场背景继续观察。"),
        ("但具体仍需结合市场背景继续观察。", "仍需结合后续市场数据与事件线索观察。"),
        ("，暂难确认具体驱动力。", "，需结合标的走势与外部事件继续观察。"),
        ("暂难确认具体驱动力。", "需结合标的走势与外部事件继续观察。"),
        ("，但具体驱动力尚不明确。", "，仍需结合后续市场数据与事件线索观察。"),
        ("但具体驱动力尚不明确。", "仍需结合后续市场数据与事件线索观察。"),
        ("具体驱动力尚不明确。", "仍需结合后续市场数据与事件线索观察。"),
        ("，但缺乏明确的外部事件或政策信号，暂作保守观察。", "，仍需结合央行操作、资金面价格与后续政策信号观察。"),
        ("但缺乏明确的外部事件或政策信号，暂作保守观察。", "仍需结合央行操作、资金面价格与后续政策信号观察。"),
        ("，但缺乏标的指数方向及外部事件证据，", "，仍需结合标的指数方向及外部事件观察，"),
        ("但缺乏标的指数方向及外部事件证据，", "仍需结合标的指数方向及外部事件观察，"),
        ("但尚无法确认是否存在增量资金入场或对冲需求。", "仍需结合后续成交、持仓与标的指数走势观察。"),
        ("，但无法确认具体原因，", "，仍需结合市场背景与后续数据观察，"),
        ("但无法确认具体原因，", "仍需结合市场背景与后续数据观察，"),
    )
    normalized = text
    for old, new in replacements:
        normalized = normalized.replace(old, new)
    normalized = normalized.strip()
    return normalized or fallback


def _sentence_end(text: str) -> str:
    stripped = text.strip().rstrip("，,；;：:")
    if not stripped:
        return ""
    if stripped[-1] in "。！？!?":
        return stripped
    return f"{stripped}。"


def _first_sentence(text: str, max_len: int = 88) -> str:
    compact = " ".join((text or "").strip().split())
    if not compact:
        return ""
    candidates = [_sentence_end(compact)]
    for sep in ("。", "；", ";"):
        if sep in compact:
            candidates.append(_sentence_end(compact.split(sep, 1)[0]))
    sentence = min(candidates, key=len)
    if len(sentence) <= max_len:
        return sentence
    for sep in ("，", ","):
        if sep in sentence:
            clause = sentence.split(sep, 1)[0]
            if clause:
                return _sentence_end(clause)
    return _sentence_end(sentence)


def _direction_label(direction: str) -> str:
    if direction == "up":
        return "上行"
    if direction == "down":
        return "下行"
    return "观察"


def _format_wow_suffix(pct: float | None) -> str:
    if pct is None:
        return ""
    sign = "+" if pct >= 0 else ""
    return f"（较上周 {sign}{_format_decimal(pct)}%）"


def _format_change_text(asset: dict[str, Any]) -> tuple[str, str]:
    field = str(asset.get("monitor_change_field", "")).strip()
    pct = _safe_float(asset.get("pct_change_pct"))
    abs_change = _safe_float(asset.get("abs_change"))
    unit = str(asset.get("unit", "")).strip()

    if field not in {"pct_change_pct", "abs_change"}:
        if pct is not None:
            field = "pct_change_pct"
        elif abs_change is not None:
            field = "abs_change"
        else:
            field = "pct_change_pct"

    if field == "abs_change":
        if abs_change is None:
            return "周度变动幅度缺失", ""
        sign = "+" if abs_change >= 0 else ""
        direction = "up" if abs_change >= 0 else "down"
        wow_suffix = _format_wow_suffix(pct)
        if _use_wan_yuan_display(asset):
            return f"周度变动 {sign}{_format_decimal(abs_change / 10000)} 万亿元{wow_suffix}", direction
        unit_suffix = f" {unit}" if unit and unit.lower() != "raw" else ""
        return f"周度变动 {sign}{_format_decimal(abs_change)}{unit_suffix}{wow_suffix}", direction

    if pct is None:
        return "周度变动幅度缺失", ""
    sign = "+" if pct >= 0 else ""
    direction = "up" if pct >= 0 else "down"
    return f"周度变动 {sign}{_format_decimal(pct)}%", direction


def _card_map(key_cards: list[Any]) -> dict[str, dict[str, Any]]:
    mapped: dict[str, dict[str, Any]] = {}
    for card in key_cards:
        if isinstance(card, dict):
            mapped[_row_join_key(card)] = card
    return mapped


def _card_texts(asset: dict[str, Any], card: dict[str, Any]) -> tuple[str, str, str]:
    direction_summary = str(card.get("direction_summary", "")).strip()
    if not direction_summary:
        direction_summary = str(card.get("card_text", "")).strip()
    direction_summary = _normalize_hs300_units(direction_summary, asset)
    magnitude_summary = _normalize_hs300_units(str(card.get("magnitude_summary", "")).strip(), asset)
    driver_analysis = _normalize_hs300_units(
        str(card.get("driver_analysis", "")).strip()
        or str(card.get("possible_drivers", "")).strip()
        or "驱动因素信息不足，暂作保守观察",
        asset,
    )
    driver_analysis = _normalize_driver_analysis_text(driver_analysis, asset)
    return (
        _normalize_decimal_places(direction_summary),
        _normalize_decimal_places(magnitude_summary),
        _normalize_decimal_places(driver_analysis),
    )


def _set_document_defaults(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 41, 55)


def _add_labeled_paragraph(doc: DocumentObject, label: str, value: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}：")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(75, 85, 99)
    _apply_run_font(label_run)
    value_run = p.add_run(value or "N/A")
    _apply_run_font(value_run)


def _add_summary_table(doc: DocumentObject, week_start: str, insight_status: str, asset_count: int) -> None:
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ["周起始", "状态", "异常资产数"]
    values = [week_start or "N/A", insight_status or "unknown", str(asset_count)]
    for idx, header in enumerate(headers):
        _set_cell_text(table.cell(0, idx), header, bold=True, size=10)
        _set_cell_text(table.cell(1, idx), values[idx], size=10)


def _set_cell_text(cell: Any, text: str, *, bold: bool = False, size: float = 9.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    _apply_run_font(run, size=size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _add_focus_section(
    doc: DocumentObject,
    top_assets: list[Any],
    cards_by_key: dict[str, dict[str, Any]],
) -> None:
    doc.add_heading("本周重点异常", level=1)
    focus_assets = [asset for asset in top_assets if isinstance(asset, dict)][:3]
    if not focus_assets:
        doc.add_paragraph("暂无重点异常。")
        return

    for idx, asset in enumerate(focus_assets, start=1):
        card = cards_by_key.get(_row_join_key(asset), {})
        asset_name = str(asset.get("asset_name", "")).strip() or "未知资产"
        metric_name = str(asset.get("metric_name", "")).strip() or "未知指标"
        move_text, _ = _format_change_text(asset)
        direction_summary, magnitude_summary, _ = _card_texts(asset, card)

        basis = _first_sentence(magnitude_summary) or _first_sentence(direction_summary) or "异常程度较高，需优先关注。"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        heading_run = p.add_run(f"{idx}. {asset_name} / {metric_name}：")
        heading_run.bold = True
        _apply_run_font(heading_run)
        body_run = p.add_run(_normalize_decimal_places(f"{move_text}；{basis}"))
        _apply_run_font(body_run)


def _add_overview_table(
    doc: DocumentObject,
    top_assets: list[Any],
    cards_by_key: dict[str, dict[str, Any]],
) -> None:
    doc.add_heading("异常资产总览", level=1)
    if not top_assets:
        doc.add_paragraph("暂无异常资产数据。")
        return

    rows = [asset for asset in top_assets if isinstance(asset, dict)][:5]
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    headers = ["排名", "资产 / 指标", "方向", "周度变动", "Agent判断依据"]
    for col_idx, header in enumerate(headers):
        _set_cell_text(table.cell(0, col_idx), header, bold=True, size=9.5)

    for row_idx, asset in enumerate(rows, start=1):
        card = cards_by_key.get(_row_join_key(asset), {})
        asset_name = str(asset.get("asset_name", "")).strip() or "未知资产"
        metric_name = str(asset.get("metric_name", "")).strip() or "未知指标"
        move_text, direction = _format_change_text(asset)
        direction_summary, magnitude_summary, _ = _card_texts(asset, card)
        basis = _first_sentence(magnitude_summary, max_len=34) or _first_sentence(direction_summary, max_len=34)
        values = [
            str(row_idx),
            f"{asset_name}\n{metric_name}",
            _direction_label(direction),
            move_text.replace("周度变动 ", ""),
            basis or "规则层识别为重点异常。",
        ]
        for col_idx, value in enumerate(values):
            _set_cell_text(table.cell(row_idx, col_idx), value, size=8.5)


def render_weekly_report_docx(report_insights: dict[str, Any]) -> DocumentObject:
    doc = Document()
    _set_document_defaults(doc)

    week_start = str(report_insights.get("week_start", "")).strip()
    top_assets = report_insights.get("top_abnormal_assets", []) or []
    key_cards = report_insights.get("key_asset_cards", []) or []
    insight_status = str(report_insights.get("insight_status", "")).strip() or "unknown"

    title = doc.add_paragraph()
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("市场异动监测周报")
    title_run.bold = True
    _apply_run_font(title_run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("仅展示 report_insights 产出的异常资产解释")
    _apply_run_font(subtitle_run)

    _add_summary_table(doc, week_start, insight_status, len(top_assets))
    doc.add_paragraph()

    if not top_assets:
        doc.add_heading("异常资产分析", level=1)
        doc.add_paragraph("暂无异常资产数据。")
        return doc

    cards_by_key = _card_map(key_cards)
    _add_focus_section(doc, top_assets, cards_by_key)
    _add_overview_table(doc, top_assets, cards_by_key)

    doc.add_page_break()
    doc.add_heading("异常资产分析", level=1)

    for idx, asset in enumerate(top_assets, start=1):
        if not isinstance(asset, dict):
            continue
        key = _row_join_key(asset)
        card = cards_by_key.get(key, {})
        asset_name = str(asset.get("asset_name", "")).strip() or "未知资产"
        metric_name = str(asset.get("metric_name", "")).strip()
        move_text, direction = _format_change_text(asset)

        heading = doc.add_heading(f"{idx}. {asset_name}", level=2)
        for run in heading.runs:
            _apply_run_font(run)
        if direction == "up":
            heading.runs[0].font.color.rgb = RGBColor(153, 27, 27)
        elif direction == "down":
            heading.runs[0].font.color.rgb = RGBColor(29, 78, 216)

        meta = doc.add_paragraph()
        label_run = meta.add_run("指标：")
        label_run.bold = True
        _apply_run_font(label_run)
        metric_run = meta.add_run(metric_name or "N/A")
        _apply_run_font(metric_run)
        spacer_run = meta.add_run("    ")
        _apply_run_font(spacer_run)
        move_run = meta.add_run(move_text)
        move_run.bold = True
        _apply_run_font(move_run)

        direction_summary, magnitude_summary, driver_analysis = _card_texts(asset, card)
        _add_labeled_paragraph(doc, "变化方向", direction_summary or "暂无方向判断")
        _add_labeled_paragraph(doc, "变化幅度", magnitude_summary or "暂无幅度判断")
        _add_labeled_paragraph(doc, "数据驱动力分析", driver_analysis)

    return doc


def write_docx_report(report_insights: dict[str, Any], output_path: str) -> None:
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = render_weekly_report_docx(report_insights)
    doc.save(path)
